from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.oxml.shared import qn, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import uuid

app = Flask(__name__)
CORS(app)
DOWNLOAD_FOLDER = "/tmp"

# =============================================================================
# NUMBERING LEVEL INDENTS (in twips) - matches Word template exactly
# =============================================================================
# Level 0 (1.):  left=720,  hanging=360
# Level 1 (a.):  left=1440, hanging=360
# Level 2 (i.):  left=2160, hanging=180
# Level 3 (1.):  left=2880, hanging=360
# Level 4 (a.):  left=3600, hanging=360

STEP_INDENTS = {
    1: {'left': 720, 'hanging': 360},
    2: {'left': 1440, 'hanging': 360},
    3: {'left': 2160, 'hanging': 180},
    4: {'left': 2880, 'hanging': 360},
    5: {'left': 3600, 'hanging': 360},
}

# Bullet indent: 360 twips left, 360 hanging (bullet at 0, text at 360)
BULLET_INDENT = {'left': 360, 'hanging': 360}


def create_numbering_definitions(doc):
    """
    Create multi-level numbering definitions in the document.
    Pattern: 1. → a. → i. → 1. → a.
    """
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        from docx.parts.numbering import NumberingPart
        numbering_part = NumberingPart.new()
        doc.part.relate_to(numbering_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
    
    numbering_xml = numbering_part._element
    
    # Check if our abstractNum already exists
    existing = numbering_xml.findall('.//w:abstractNum[@w:abstractNumId="100"]', 
                                      namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if existing:
        return
    
    # Create abstractNum for multi-level steps
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), '100')
    
    level_configs = [
        {'ilvl': '0', 'numFmt': 'decimal', 'lvlText': '%1.', 'lvlJc': 'left', 'left': 720, 'hanging': 360},
        {'ilvl': '1', 'numFmt': 'lowerLetter', 'lvlText': '%2.', 'lvlJc': 'left', 'left': 1440, 'hanging': 360},
        {'ilvl': '2', 'numFmt': 'lowerRoman', 'lvlText': '%3.', 'lvlJc': 'right', 'left': 2160, 'hanging': 180},
        {'ilvl': '3', 'numFmt': 'decimal', 'lvlText': '%4.', 'lvlJc': 'left', 'left': 2880, 'hanging': 360},
        {'ilvl': '4', 'numFmt': 'lowerLetter', 'lvlText': '%5.', 'lvlJc': 'left', 'left': 3600, 'hanging': 360},
    ]
    
    for cfg in level_configs:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), cfg['ilvl'])
        
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), cfg['numFmt'])
        lvl.append(numFmt)
        
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), cfg['lvlText'])
        lvl.append(lvlText)
        
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), cfg['lvlJc'])
        lvl.append(lvlJc)
        
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(cfg['left']))
        ind.set(qn('w:hanging'), str(cfg['hanging']))
        pPr.append(ind)
        lvl.append(pPr)
        
        abstract_num.append(lvl)
    
    first_num = numbering_xml.find('.//w:num', 
                                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_xml.append(abstract_num)
    
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '100')
    abstract_num_id = OxmlElement('w:abstractNumId')
    abstract_num_id.set(qn('w:val'), '100')
    num.append(abstract_num_id)
    numbering_xml.append(num)


def add_horizontal_rule(doc):
    """Add a horizontal rule as an empty paragraph with a bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_empty_paragraph(doc):
    """Add a blank paragraph for visual spacing."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    return para


def add_text_paragraph(doc, text, bold=False, size=11):
    """Add a simple text paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Calibri'
    return para


def add_labelled_paragraph(doc, label, value):
    """Add a labelled paragraph with bold label and normal value (inline)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    
    run1 = para.add_run(f"{label}: ")
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.name = 'Calibri'
    
    if value:
        run2 = para.add_run(value)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0, 0, 0)
        run2.font.name = 'Calibri'
    
    return para


def add_label_only(doc, label):
    """Add just a label with colon (for when bullets follow on next lines)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(f"{label}:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Calibri'
    return para


def add_bullet(doc, text, indent_level=0):
    """
    Add a bullet point with proper hanging indent so wrapped text aligns.
    Uses left indent + negative first-line (hanging) so text wraps properly.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    
    # Hanging indent: left margin where text wraps, first-line negative pulls bullet back
    # Level 0: bullet at 360, text at 720
    # Level 1: bullet at 1080, text at 1440
    left_twips = 720 + (indent_level * 720)
    hanging_twips = 360
    
    para.paragraph_format.left_indent = Twips(left_twips)
    para.paragraph_format.first_line_indent = Twips(-hanging_twips)  # Negative = hanging
    
    # Add bullet character (bold) and text (normal)
    bullet_run = para.add_run("• ")
    bullet_run.bold = True
    bullet_run.font.size = Pt(11)
    bullet_run.font.color.rgb = RGBColor(0, 0, 0)
    bullet_run.font.name = 'Calibri'
    
    text_run = para.add_run(text)
    text_run.font.size = Pt(11)
    text_run.font.color.rgb = RGBColor(0, 0, 0)
    text_run.font.name = 'Calibri'
    
    return para


def add_numbered_step(doc, text, level):
    """
    Add a numbered step using Word's multi-level numbering.
    Level 1 = 1., Level 2 = a., Level 3 = i., Level 4 = 1., Level 5 = a.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    
    # Apply ListParagraph style if it exists
    try:
        para.style = doc.styles['List Paragraph']
    except:
        pass
    
    # Apply numbering via XML
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level - 1))  # Convert 1-5 to 0-4
    numPr.append(ilvl)
    
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '100')
    numPr.append(numId)
    
    pPr.insert(0, numPr)
    
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Calibri'
    return para


def add_note(doc, text, preceding_level):
    """
    Add an italic note paragraph aligned with the preceding step level.
    The note's left indent matches the left indent of that step level.
    """
    # Get the left indent for the preceding level (in twips)
    level_indent = STEP_INDENTS.get(preceding_level, STEP_INDENTS[5])
    left_twips = level_indent['left']
    
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.left_indent = Twips(left_twips)
    
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Calibri'
    return para


def add_revision_table(doc, rows_data):
    """
    Add the Revision History table matching template format:
    - Borderless table (no outer borders, no inside borders)
    - Bold headers with bottom border (sz=12)
    - First data row has top border (sz=12) creating single separator line
    - Column widths as percentage: ~27.4%, ~19.7%, ~52.8%
    """
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Set table to 100% width using percentage type
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # Table width 5000 = 100% in pct type
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    
    # Remove all default borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Column widths in pct (out of 5000): 1372, 987, 2641
    col_widths_pct = [1372, 987, 2641]
    headers = ["Date", "Revised By", "Description"]
    
    # Header row
    header_cells = table.rows[0].cells
    for i, (cell, header, width_pct) in enumerate(zip(header_cells, headers, col_widths_pct)):
        # Set cell width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_pct))
        tcW.set(qn('w:type'), 'pct')
        tcPr.append(tcW)
        
        # Add bottom border to header cells
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5pt
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)
        
        # Add header text
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Calibri'
    
    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        if isinstance(row_data, dict):
            if 'text' in row_data:
                parts = row_data['text'].split('|||')
                values = [p.strip() for p in parts]
            else:
                values = [
                    row_data.get('date', ''),
                    row_data.get('revised_by', ''),
                    row_data.get('description', '')
                ]
        elif isinstance(row_data, str):
            parts = row_data.split('|||')
            values = [p.strip() for p in parts]
        else:
            values = ['', '', '']
        
        # Ensure we have 3 values
        while len(values) < 3:
            values.append('')
        
        row = table.add_row()
        for i, (cell, value, width_pct) in enumerate(zip(row.cells, values[:3], col_widths_pct)):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Set cell width
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width_pct))
            tcW.set(qn('w:type'), 'pct')
            tcPr.append(tcW)
            
            # Add top border to first data row cells (creates the separator line)
            if row_idx == 0:
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '12')
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), 'auto')
                tcBorders.append(top)
                tcPr.append(tcBorders)
            
            # Add cell text
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Calibri'
    
    return table


def setup_footer(doc, sop_title, sop_id, revision_date):
    """Set up footer. First page blank, subsequent pages have SOP info centered."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    # Default footer (pages 2+)
    footer = section.footer
    footer.is_linked_to_previous = False
    
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    # Blank line
    para1 = footer.add_paragraph()
    para1.paragraph_format.space_after = Pt(0)
    
    # SOP title line
    para2 = footer.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_after = Pt(0)
    run2 = para2.add_run(f"{sop_title} [{sop_id}]")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0, 0, 0)
    run2.font.name = 'Calibri'
    
    # Revision date line
    para3 = footer.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para3.paragraph_format.space_after = Pt(0)
    run3 = para3.add_run(f"Revision Date: {revision_date}")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0, 0, 0)
    run3.font.name = 'Calibri'
    
    # First page footer (blank)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    for para in first_footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
    first_footer.add_paragraph()


# Labels that trigger a blank line before the NEXT different label in same section
LABELS_NEEDING_SPACE_AFTER = {
    'Objective', 'Objectives',
    'Process Owner', 'Process Owners', 'Process Owner(s)',
    'Input', 'Inputs', 'Input(s)',
    'Dependency', 'Dependencies', 'Dependency(ies)',
}

# Normalize label names for comparison
def normalize_label(label):
    """Normalize label for comparison (remove trailing s, parens, etc.)"""
    label = label.strip()
    # Remove trailing (s) or (ies)
    if label.endswith('(s)'):
        label = label[:-3]
    if label.endswith('(ies)'):
        label = label[:-5] + 'y'
    # Remove trailing s for plurals
    if label.endswith('s') and not label.endswith('ss'):
        label = label[:-1]
    return label.lower()


def generate_sop_doc(data):
    """
    Generate an SOP document from the provided data.
    """
    try:
        doc = Document("template.docx")
        for para in doc.paragraphs[:]:
            p = para._element
            p.getparent().remove(p)
    except:
        doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.footer_distance = Inches(0.5)
    section.header_distance = Inches(0.5)
    
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Create numbering definitions
    create_numbering_definitions(doc)
    
    # Header data
    sop_title = data.get('title', 'Generated SOP')
    sop_id = data.get('sop_id', '') or 'SOP-ID TBD'
    prepared_by = data.get('prepared_by', '')
    approved_by = data.get('approved_by', '')
    revision_date = data.get('revision_date', '')
    
    # === HEADER BLOCK ===
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(f"SOP Title: {sop_title}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Calibri'
    
    add_text_paragraph(doc, f"SOP ID: {sop_id}")
    add_empty_paragraph(doc)
    add_text_paragraph(doc, f"Prepared By: {prepared_by}")
    add_text_paragraph(doc, f"Approved By: {approved_by}")
    add_text_paragraph(doc, f"Revision Date: {revision_date}")
    add_horizontal_rule(doc)
    
    # === SECTIONS ===
    sections_data = data.get('sections', [])
    last_step_level = 1  # Track for note indentation
    
    for sec_idx, sec_data in enumerate(sections_data):
        heading = sec_data.get('heading', '')
        
        if heading:
            add_text_paragraph(doc, heading, bold=True)
            add_empty_paragraph(doc)
        
        # Handle table type (Revision History)
        if sec_data.get('type') == 'table':
            content = sec_data.get('content', [])
            add_revision_table(doc, content)
        else:
            content = sec_data.get('content', [])
            i = 0
            last_label = None
            had_bullets_after_label = False
            
            while i < len(content):
                item = content[i]
                if not isinstance(item, dict):
                    i += 1
                    continue
                
                item_type = item.get('type', 'text')
                text = item.get('text', '').strip()
                
                if not text and item_type != 'spacer':
                    i += 1
                    continue
                
                if item_type == 'heading':
                    add_text_paragraph(doc, text, bold=True)
                    add_empty_paragraph(doc)
                    last_label = None
                    had_bullets_after_label = False
                
                elif item_type == 'labelled':
                    # Extract label from text
                    if ':' in text:
                        new_label, _, value = text.partition(':')
                        new_label = new_label.strip()
                        value = value.strip()
                    else:
                        new_label = text
                        value = ''
                    
                    # Check if we need spacing before this label
                    # (if previous label had bullets and this is a different label)
                    if last_label and had_bullets_after_label:
                        norm_last = normalize_label(last_label)
                        norm_new = normalize_label(new_label)
                        if norm_last != norm_new:
                            # Check if last label is one that needs space after
                            for trigger_label in LABELS_NEEDING_SPACE_AFTER:
                                if normalize_label(trigger_label) == norm_last:
                                    add_empty_paragraph(doc)
                                    break
                    
                    # Check if bullets follow this label
                    bullets_follow = False
                    for j in range(i + 1, len(content)):
                        next_item = content[j]
                        if isinstance(next_item, dict):
                            next_type = next_item.get('type', '')
                            if next_type in ('bullet', 'sub_bullet'):
                                bullets_follow = True
                                break
                            elif next_type not in ('spacer',):
                                break
                    
                    if ':' in text:
                        if bullets_follow or not value:
                            add_label_only(doc, new_label)
                        else:
                            add_labelled_paragraph(doc, new_label, value)
                    else:
                        add_text_paragraph(doc, text)
                    
                    last_label = new_label
                    had_bullets_after_label = False
                
                elif item_type == 'bullet':
                    add_bullet(doc, text, indent_level=0)
                    had_bullets_after_label = True
                
                elif item_type == 'sub_bullet':
                    add_bullet(doc, text, indent_level=1)
                    had_bullets_after_label = True
                
                elif item_type == 'step':
                    level = item.get('level', 1)
                    level = max(1, min(5, level))
                    add_numbered_step(doc, text, level)
                    last_step_level = level
                    last_label = None
                    had_bullets_after_label = False
                
                elif item_type == 'note':
                    add_empty_paragraph(doc)
                    add_note(doc, text, last_step_level)
                    add_empty_paragraph(doc)
                
                elif item_type == 'spacer':
                    add_empty_paragraph(doc)
                
                else:
                    add_text_paragraph(doc, text)
                
                i += 1
        
        # Horizontal rule between sections (not after last)
        if sec_idx < len(sections_data) - 1:
            add_horizontal_rule(doc)
    
    # Footer
    setup_footer(doc, sop_title, sop_id, revision_date)
    
    # Save
    filename = f"sop_{uuid.uuid4().hex}.docx"
    filepath = os.path.join(DOWNLOAD_FOLDER, filename)
    doc.save(filepath)
    return filename


@app.route("/download/<filename>", methods=["GET"])
def download_file(filename):
    filepath = os.path.join(DOWNLOAD_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    return {"error": "File not found."}, 404


@app.route("/generate", methods=["POST"])
def generate():
    try:
        data = request.json
        filename = generate_sop_doc(data)
        link = f"https://sop-flask-api.onrender.com/download/{filename}"
        return jsonify({"download_link": link})
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}, 500


@app.route("/health", methods=["GET"])
def health():
    return {"status": "ok"}


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080)
